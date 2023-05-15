import tkinter
import tkinter.messagebox
import customtkinter
import os, shutil, requests, datetime, time

from threading import Thread
from ics import Calendar
from dateutil.relativedelta import relativedelta
from openpyxl import load_workbook
from tkinter import filedialog
from PIL import Image, ExifTags
from docxtpl import DocxTemplate
from docx import Document
from docx2pdf import convert
from docx.shared import Cm
from dotenv import load_dotenv


basedir = os.path.dirname(__file__)
load_dotenv()   # 使用dotenv保存日历链接
calender_url = os.environ.get('CALENDAR_URL')
customtkinter.set_appearance_mode("Light")  # Modes: "System" (standard), "Dark", "Light"
customtkinter.set_default_color_theme("blue")  # Themes: "blue" (standard), "green", "dark-blue"


class info_window(customtkinter.CTkToplevel):
    def __init__(self, info, progress_bar=False, *args, **kwargs):
        super().__init__(*args, **kwargs)
        # self.geometry("400*100")

        window_width = 400
        window_height = 100
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()-100
        x_cordinate = int((screen_width/2) - (window_width/2))
        y_cordinate = int((screen_height/2) - (window_height/2))
        # self.geometry(f"{1000}x{580}")
        self.geometry("{}x{}+{}+{}".format(window_width, window_height, x_cordinate, y_cordinate))
        self.attributes('-topmost', 'true')

        normal_font = customtkinter.CTkFont(family="微软雅黑", size=13, weight="bold")
        
        if progress_bar:
            self.title('处理中')
            # info = '请稍后'
            self.label = customtkinter.CTkLabel(self, text=info, font=normal_font)
            self.label.grid(row=0, column=0, padx=(50,50), pady=(30,5))
            self.progressbar = customtkinter.CTkProgressBar(self)
            self.progressbar.grid(row=1, column=0, padx=(100, 100), pady=(5, 50), sticky="ew")
            self.progressbar.configure(mode="indeterminnate")
            self.progressbar.start()
            
        else:
            self.title('错误')
            self.label = customtkinter.CTkLabel(self, text=info, font=normal_font)
            self.label.pack(padx=(50,50), pady=(30,30))
        



class App(customtkinter.CTk):
    def __init__(self):
        super().__init__()

        # title_font = customtkinter.CTkFont(family="微软雅黑", size=20, weight="bold")
        normal_font = customtkinter.CTkFont(family="微软雅黑", size=13, weight="bold")
        self.toplevel_window = None
        
        # configure window
        self.title("Future X Baotou")
        window_width = 1000
        window_height = 580
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()-100
        x_cordinate = int((screen_width/2) - (window_width/2))
        y_cordinate = int((screen_height/2) - (window_height/2))
        # self.geometry(f"{1000}x{580}")
        self.geometry("{}x{}+{}+{}".format(window_width, window_height, x_cordinate, y_cordinate))

        # configure grid layout (4x4)
        self.grid_columnconfigure(1, weight=1)
        self.grid_columnconfigure((2, 3), weight=0)
        self.grid_rowconfigure((0, 1, 2), weight=1)

        # 左侧边栏
        self.sidebar_frame = customtkinter.CTkFrame(self, width=140, corner_radius=0)
        self.sidebar_frame.grid(row=0, column=0, rowspan=6, sticky="nsew")
        self.sidebar_frame.grid_rowconfigure(4, weight=1)

        self.appearance_mode_label = customtkinter.CTkLabel(self.sidebar_frame, text="课程类别:", font=normal_font, anchor="w")
        self.appearance_mode_label.grid(row=0, column=0, padx=20, pady=(20, 0))

        self.appearance_mode_optionemenu = customtkinter.CTkOptionMenu(self.sidebar_frame, values=["Light", "Dark", "System"], command=self.change_appearance_mode_event)
        self.appearance_mode_optionemenu.grid(row=1, column=0, padx=20, pady=(5, 10))

        self.scaling_label = customtkinter.CTkLabel(self.sidebar_frame, text="课程选择:", font=normal_font, anchor="w")
        self.scaling_label.grid(row=2, column=0, padx=20, pady=(10, 0))
        
        self.scaling_optionemenu = customtkinter.CTkOptionMenu(self.sidebar_frame, values=["80%", "90%", "100%", "110%", "120%"], command=self.change_scaling_event)
        self.scaling_optionemenu.grid(row=3, column=0, padx=20, pady=(5, 20))

        self.sidebar_button_generate_report = customtkinter.CTkButton(self.sidebar_frame, command=self.genereate_report_event)
        self.sidebar_button_generate_report.grid(row=6, column=0, padx=20, pady=10)

        self.sidebar_button_convert_2_PDF = customtkinter.CTkButton(self.sidebar_frame, command=self.convert_2_PDF_event)
        self.sidebar_button_convert_2_PDF.grid(row=7, column=0, padx=20, pady=10)

        self.sidebar_button_statistics = customtkinter.CTkButton(self.sidebar_frame, command=self.statistics_event)
        self.sidebar_button_statistics.grid(row=8, column=0, padx=20, pady=10)

        # self.sidebar_button_3 = customtkinter.CTkButton(self.sidebar_frame, command=self.sidebar_button_event)
        # self.sidebar_button_3.grid(row=3, column=0, padx=20, pady=10)
        


        # 主文本框
        self.textbox_lecture_comment = customtkinter.CTkTextbox(self, width=500)
        self.textbox_lecture_comment.grid(row=0, column=1, columnspan=3, padx=(20, 20), pady=(20, 0), sticky="ew")


        # 文本框下方内容
        self.slider_progressbar_frame = customtkinter.CTkFrame(self, fg_color="transparent")
        self.slider_progressbar_frame.grid(row=1, column=1, padx=10, pady=(10, 0), sticky="nsew")
        self.slider_progressbar_frame.grid_columnconfigure((0,1,2), weight=1)
        self.slider_progressbar_frame.grid_rowconfigure(4, weight=1)

        self.textbox_1_label = customtkinter.CTkLabel(self.slider_progressbar_frame, text="教师姓名:", font=normal_font, anchor="w")
        self.textbox_1_label.grid(row=0, column=0, padx=10, pady=(10, 2), sticky='w')
        self.textbox_teacher_name = customtkinter.CTkTextbox(self.slider_progressbar_frame, height=20, width=55)
        self.textbox_teacher_name.grid(row=1, column=0, padx=(10, 10), pady=(0, 5), sticky="ew")

        self.textbox_2_label = customtkinter.CTkLabel(self.slider_progressbar_frame, text="学生姓名:", font=normal_font, anchor="w")
        self.textbox_2_label.grid(row=0, column=1, padx=10, pady=(10, 2), sticky='w')
        self.textbox_student_name = customtkinter.CTkTextbox(self.slider_progressbar_frame, height=20, width=60)
        self.textbox_student_name.grid(row=1, column=1, padx=(10, 10), pady=(0, 5), sticky="ew")

        

        self.seg_button_1_label = customtkinter.CTkLabel(self.slider_progressbar_frame, text="思辨与交流能力:", font=normal_font, anchor="w")
        self.seg_button_1_label.grid(row=2, column=0, padx=10, pady=(10, 2), sticky='w')
        self.seg_button_1 = customtkinter.CTkSegmentedButton(self.slider_progressbar_frame)
        self.seg_button_1.grid(row=3, column=0, padx=(10, 10), pady=(0, 5), sticky="ew")

        self.seg_button_2_label = customtkinter.CTkLabel(self.slider_progressbar_frame, text="反思与创新能力:", font=normal_font, anchor="w")
        self.seg_button_2_label.grid(row=2, column=1, padx=10, pady=(10, 2), sticky='w')
        self.seg_button_2 = customtkinter.CTkSegmentedButton(self.slider_progressbar_frame)
        self.seg_button_2.grid(row=3, column=1, padx=(10, 10), pady=(0, 5), sticky="ew")
        
        self.seg_button_3_label = customtkinter.CTkLabel(self.slider_progressbar_frame, text="合作与互助能力:", font=normal_font, anchor="w")
        self.seg_button_3_label.grid(row=2, column=2, padx=10, pady=(10, 2), sticky='w')
        self.seg_button_3 = customtkinter.CTkSegmentedButton(self.slider_progressbar_frame)
        self.seg_button_3.grid(row=3, column=2, padx=(10, 10), pady=(0, 5), sticky="ew")

        self.seg_button_4_label = customtkinter.CTkLabel(self.slider_progressbar_frame, text="问题解决能力:", font=normal_font, anchor="w")
        self.seg_button_4_label.grid(row=4, column=0, padx=10, pady=(5, 2), sticky='w')
        self.seg_button_4 = customtkinter.CTkSegmentedButton(self.slider_progressbar_frame)
        self.seg_button_4.grid(row=5, column=0, padx=(10, 10), pady=(0, 5), sticky="ew")

        self.seg_button_5_label = customtkinter.CTkLabel(self.slider_progressbar_frame, text="编程思维:", font=normal_font, anchor="w")
        self.seg_button_5_label.grid(row=4, column=1, padx=10, pady=(5, 2), sticky='w')
        self.seg_button_5 = customtkinter.CTkSegmentedButton(self.slider_progressbar_frame)
        self.seg_button_5.grid(row=5, column=1, padx=(10, 10), pady=(0, 5), sticky="ew")


        # 下方输入栏
        # 选择课评模板
        self.entry_template = customtkinter.CTkEntry(self,)
        self.entry_template.grid(row=3, column=1, columnspan=2, padx=(20, 0), pady=(20, 2), sticky="nsew")
        self.button_select_template = customtkinter.CTkButton(master=self, fg_color="transparent", border_width=2, text_color=("gray10", "#DCE4EE"), command=self.select_template_event)
        self.button_select_template.grid(row=3, column=3, padx=(20, 20), pady=(20, 2), sticky="nsew")
        # 选择照片文件
        self.entry_1 = customtkinter.CTkEntry(self,)
        self.entry_1.grid(row=4, column=1, columnspan=2, padx=(20, 0), pady=(20, 2), sticky="nsew")
        self.button_select_photos = customtkinter.CTkButton(master=self, fg_color="transparent", border_width=2, text_color=("gray10", "#DCE4EE"), command=self.select_photo_event)
        self.button_select_photos.grid(row=4, column=3, padx=(20, 20), pady=(20, 2), sticky="nsew")
        #选择保存地址
        self.entry_2 = customtkinter.CTkEntry(self)
        self.entry_2.grid(row=5, column=1, columnspan=2, padx=(20, 0), pady=(2, 20), sticky="nsew")
        self.button_saving_location = customtkinter.CTkButton(master=self, fg_color="transparent", border_width=2, text_color=("gray10", "#DCE4EE"), command=self.saving_location_event)
        self.button_saving_location.grid(row=5, column=3, padx=(20, 20), pady=(2, 20), sticky="nsew")


        # 设置默认数值
        # self.sidebar_button_3.configure(state="disabled", text="Disabled CTkButton")
        self.sidebar_button_generate_report.configure(text="生成报告", font=normal_font)
        self.sidebar_button_convert_2_PDF.configure(text="PDF转换", font=normal_font)
        self.sidebar_button_statistics.configure(text="课时统计", font=normal_font)
        self.button_select_template.configure(text="选择模板", font=normal_font)
        self.button_select_photos.configure(text="选择照片", font=normal_font)
        self.button_saving_location.configure(text="选择保存位置", font=normal_font)
        self.appearance_mode_optionemenu.set("Dark")
        self.scaling_optionemenu.set("100%")
        self.textbox_lecture_comment.insert("0.0", "在此处输入教师评价")
        self.seg_button_1.configure(values=["1","2","3","4","5"], font=normal_font)
        self.seg_button_1.set("3")
        self.seg_button_2.configure(values=["1","2","3","4","5"], font=normal_font)
        self.seg_button_2.set("3")
        self.seg_button_3.configure(values=["1","2","3","4","5"], font=normal_font)
        self.seg_button_3.set("3")
        self.seg_button_4.configure(values=["1","2","3","4","5"], font=normal_font)
        self.seg_button_4.set("3")
        self.seg_button_5.configure(values=["1","2","3","4","5"], font=normal_font)
        self.seg_button_5.set("3")

    # 按钮功能
    # def open_input_dialog_event(self):
    #     dialog = customtkinter.CTkInputDialog(text="Type in a number:", title="CTkInputDialog")
    #     print("CTkInputDialog:", dialog.get_input())

    def genereate_report_event(self):   # 生成课程评价，并添加照片
        dir_path = os.path.dirname(os.path.realpath(__file__))

        pict = Image.open(f_path[0])
        exif_data = pict._getexif()
        picDate = exif_data[36867]
        pict.close()

        # name = docList[0]
        name = doc_path
        communication = int(self.seg_button_1.get())
        creation = int(self.seg_button_2.get())
        co_operation = int(self.seg_button_3.get())
        solvability =int(self.seg_button_4.get())
        thoughts =int(self.seg_button_5.get())
        lecturer_comments = self.textbox_lecture_comment.get("1.0", "end-1c")
        global teacher_name
        teacher_name = self.textbox_teacher_name.get("1.0", "end-1c")
        global student_name
        student_name = self.textbox_student_name.get("1.0", "end-1c")
        year = picDate[0:4]
        month = picDate[5:7]
        day = picDate[8:10]

        saving_folder = saving_location+'/'+student_name
        number = 0
        while True: # 检测保存目录是否存在，存在则在文件夹后增加数字
            if not os.path.exists(saving_folder):
                os.makedirs(saving_folder)
                break
            else:
                number += 1
                saving_folder = saving_folder + str(number)

        doc = DocxTemplate(name) #加载模板文件
        document = Document(name)

        data_dic = {
            'teacher_name' : teacher_name,
            'student_name' : student_name,
            'communication' : '★'*(communication-1),
            'creation' : '★'*(creation-1),
            'co_operation' : '★'*(co_operation-1),
            'solvability' : '★'*(solvability-1),
            'thoughts' : '★'*(thoughts-1),
            'lecturer_comments' : lecturer_comments,
            'year' : year,
            'month' : month,
            'day' : day
        }
        doc.render(data_dic) #填充数据
        table = doc.tables[0] # 获取课程名称
        global lessonName
        lessonName = table.cell(1, 1).text
        documentName = student_name+lessonName+str(year)+str(month)+str(day)+'课评报告.docx'
        doc.save(documentName) #保存目标文件

        docAddPicName = documentName
        docAddPic = Document(docAddPicName)
        paragraph = docAddPic.add_paragraph()

        # 检测图片是否需要旋转
        n = 0
        for picName in f_path:
            try:
                image=Image.open(picName)
                for orientation in ExifTags.TAGS.keys():
                    if ExifTags.TAGS[orientation]=='Orientation':
                        break
                exif = image._getexif()
                try:
                    if exif[orientation] == 3:
                        image=image.rotate(180, expand=True)
                    elif exif[orientation] == 6:
                        image=image.rotate(270, expand=True)
                    elif exif[orientation] == 8:
                        image=image.rotate(90, expand=True)
                    # print('process done!')
                except:
                    # print('no need to process!')
                    pass
                image.save(picName)
                image.close()

                # 添加照片到文件中
                run = paragraph.add_run()
                run.add_picture(picName, width=Cm(6))
                docAddPic.save(documentName)
                print('photo added!')
                image.close()

                # 更改照片名
                picNameStandard = student_name+year+month+day+'_'+str(n)+'.jpg'
                os.rename(picName, picNameStandard)
                n = n+1

                # 移动照片与课评文档至保存路径中
                shutil.move(dir_path+'/'+picNameStandard, saving_folder+'/'+picNameStandard)
                shutil.move(dir_path+'/'+documentName, saving_folder+'/'+documentName)
                
                info = '已生成' + ' ' + student_name + '_' + lessonName + ' ' + '课评报告'
                self.pass_info_2_top_level(info=info)

            except (AttributeError, KeyError, IndexError):
                # cases: image don't have getexif
                print('no need to process!')

    def select_template_event(self):    # 选择课程评价模版
        self.entry_template.delete("0", "end")
        root = tkinter.Tk()
        root.withdraw()
        global doc_path
        # os.startfile(str(os.getcwd())+'\\assest\\课程评价')
        doc_path = filedialog.askopenfilename(initialdir=(basedir+'/assest/课程评价'))
        self.entry_template.insert("0", doc_path)

    def select_photo_event(self):   # 选择课程评价中的照片
        self.entry_1.delete("0", "end")
        root = tkinter.Tk()
        root.withdraw()
        global f_path 
        f_path = filedialog.askopenfilenames()
        print('\n获取的文件地址：', f_path)
        photo_name_string = ''
        for photo in f_path:
            photo_path = os.path.split(photo)
            photo_name_string += (photo_path[1]+'  ')
        print(photo_name_string)
        self.entry_1.insert("0", photo_name_string)

    def saving_location_event(self):    # 使用tk窗口打开保存目录
        self.entry_2.delete("0", "end")
        root = tkinter.Tk()
        root.withdraw()
        global saving_location
        saving_location = filedialog.askdirectory()
        self.entry_2.insert("0", saving_location)
        # 使用保存目录按钮测试弹出窗口
        # info = 'wait'
        # if self.toplevel_window is None or not self.toplevel_window.winfo_exists():
        #     self.toplevel_window = self.pass_info_2_top_level(info=info,progress_bar=True)  # create window if its None or destroyed
        # else:
        #     self.toplevel_window.focus()  # if window exists focus it
        

    def convert_2_PDF_event(self):
        info = '正在转换，请稍后'
        self.pass_info_2_top_level(info=info, progress_bar=True)
        
        global t1
        t1 = Thread(target=self.convert_2_PDF)
        t1.start()

    def convert_2_PDF(self):  # 将保存目录中的word文件全部转换为pdf文件
        try:
            shutil.rmtree('temp')
        except:
            os.mkdir('temp')

        try:
            folder_path = saving_location
        except NameError:
            self.toplevel_window.destroy()
            self.pass_info_2_top_level(info='未选择保存位置')
        file_list = []

        def CrossOver(dir, file_list):
            try:
                for i in os.listdir(dir):  # 遍历整个文件夹
                    path = os.path.join(dir, i)
                    if os.path.isfile(path):  # 判断是否为一个文件，排除文件夹
                        if os.path.splitext(path)[1]==".docx":  # 判断文件扩展名是否为“.docx”
                            temp_path = 'temp/' + os.path.split(path)[1]
                            shutil.copy2(path, temp_path)
                            file_list.append(path)
                    elif os.path.isdir(path):
                        newdir=path
                        CrossOver(newdir, file_list)
                return len(file_list)
            except FileNotFoundError:
                self.toplevel_window.destroy()
                self.pass_info_2_top_level(info='未选择保存位置')


        output = CrossOver(folder_path, file_list)   # 执行函数，输出结果
        print('已获取' + folder_path + '中的' + str(output) + '个文件')
        
        convert('temp/')    # 批量转换保存目录中的word文档

        count = 0
        for old_file in os.listdir('temp'):
            if os.path.split(old_file)[1].endswith('.pdf'):
                count += 1
                new_file = saving_location + '/' + os.path.split(old_file)[1]
                shutil.copyfile('temp/'+old_file, new_file)
        shutil.rmtree('temp')

        info = '已转换' + str(count) + '个文件为PDF'
        print(info)
        self.toplevel_window.destroy()
        time.sleep(0.5)
        self.pass_info_2_top_level(info=info)

    def statistics_event(self):     # 课时统计按钮，弹出窗口并输入月份
        dialog = customtkinter.CTkInputDialog(text="输入需要导出的月份:", title="课时统计")
        month = int(dialog.get_input())
        self.statistics(month)
        # print('正在导出', dialog.get_input(),'月课时统计')
    
    def statistics(self, month):     # 统计本月的教师课时
        c = Calendar(requests.get(calender_url).text)

        now = datetime.datetime.now()
        year = now.year
        month = month

        start_date = datetime.date(year,month,1)
        end_date = start_date + relativedelta(months=1)

        new_file_name = '图图'+str(year)+'年'+str(month)+'月课时统计.xlsx'

        e = list(c.timeline)    # 日历中的全部事件

        class_in_total = []
        for calender_event in e:    # 将事件添加到class_in_total中
            current_class = []
            if calender_event.begin.date() >= start_date and calender_event.begin.date() < end_date: # 筛选指定日期内的事件
                try:
                    c1ass = calender_event.name.split(' ')[0]
                    student = len(calender_event.name.split(' ')[1:])
                    class_date = calender_event.begin.format('YYYY-MM-DD HH:mm').split(' ')[0]
                    class_time = calender_event.begin.format('YYYY-MM-DD HH:mm').split(' ')[1]
                    current_class.append(c1ass)
                    current_class.append(class_date)
                    current_class.append(class_time)
                    current_class.append(student)
                except IndexError:
                    pass

                class_in_total.append(current_class)
        class_in_total.sort()
        print('日历中共有'+str(len(class_in_total))+'项日程')

        wb = load_workbook(filename = 'assest/教师课时统计表模板.xlsx')
        ws = wb.active

        cols = {'启蒙':2, '玛塔':5, '程小奔':5, '探究':8, '工程':11, 'Scratch':14, 'Python':14, '南开':17, '普林斯顿':17, '博苑澳森':17, '考古营':24, '医学营':24}

        def insert_course_information(calender_list_item, course):  # 将内容添加至excel中
            time_value = datetime.datetime.strptime(calender_list_item[1],'%Y-%m-%d')   # 整理时间格式
            time_value = time_value.strftime('%m月%d日')
            if calender_list_item[0] == course:
                row = 5     # 固定在第5行开始
                if course not in cols.keys():   # 跳过日历中的其他事项
                    print('未添加 ',time_value,course)
                else:
                    col = cols[course]  # 从cols中获取课程所在列

                    while ws.cell(row=row, column=col).value != None :  # 判断该单元格是否为空
                        # print(str(row) + ' ' +str(col) + ' is not empty')
                        row += 1
                        if ws.cell(row=row, column=col).value == None:
                            break

                    if course in ('南开','普林斯顿','博苑澳森'):    # 课程格式不同
                        ws.cell(row=row, column=col, value=course)
                        ws.cell(row=row, column=col+1, value=time_value)
                        ws.cell(row=row, column=col+2, value=calender_list_item[2])
                        ws.cell(row=row, column=col+3, value=calender_list_item[3])
                        print('已添加',time_value,course,calender_list_item[3],'人')
                        row += 1
                        wb.save(filename='教师课时统计表模板.xlsx')
                    else:
                        ws.cell(row=row, column=col, value=time_value)
                        ws.cell(row=row, column=col+1, value=calender_list_item[2])
                        ws.cell(row=row, column=col+2, value=calender_list_item[3])
                        print('已添加',time_value,course,calender_list_item[3],'人')
                        row += 1
                        wb.save(filename='教师课时统计表模板.xlsx')

        for item in class_in_total:
            insert_course_information(item, item[0])
        print('已导入'+str(len(class_in_total))+'项日程')

        os.rename('教师课时统计表模板.xlsx', new_file_name)
        dir_path = os.path.dirname(os.path.realpath(__file__))
        try:
            shutil.move(dir_path+'/'+new_file_name, saving_location+'/'+new_file_name)
            self.pass_info_2_top_level(info='已生成'+' '+new_file_name)
        except NameError:
            self.pass_info_2_top_level(info='未选择保存目录')
            os.remove(new_file_name)


    def pass_info_2_top_level(self, info, progress_bar=False):
        if self.toplevel_window is None or not self.toplevel_window.winfo_exists():
            self.toplevel_window = info_window(info=info, progress_bar=progress_bar)  # create window if its None or destroyed
        else:
            self.toplevel_window.focus()  # if window exists focus it


    def change_appearance_mode_event(self, new_appearance_mode: str):
        customtkinter.set_appearance_mode(new_appearance_mode)

    def change_scaling_event(self, new_scaling: str):
        new_scaling_float = int(new_scaling.replace("%", "")) / 100
        customtkinter.set_widget_scaling(new_scaling_float)

    

if __name__ == "__main__":
    # root.protocol("WM_DELETE_WINDOW", _quit)
    app = App()
    def quit():
        app.quit()
        app.destroy()

    app.protocol('WM_DELETE_WINDOW', quit)
    app.mainloop()

    
    
    